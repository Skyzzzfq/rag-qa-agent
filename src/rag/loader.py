import os
from pathlib import Path

from langchain_core.documents import Document
from PyPDF2 import PdfReader

from src.utils.logger import get_logger

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".md", ".markdown", ".txt", ".docx"}


def load_pdf(file_path: str) -> list[Document]:
    """加载 PDF 文件，每页生成一个 Document"""
    documents: list[Document] = []
    try:
        reader = PdfReader(file_path)
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                documents.append(
                    Document(
                        page_content=text.strip(),
                        metadata={
                            "source": os.path.basename(file_path),
                            "page": page_num,
                            "type": "pdf",
                        },
                    )
                )
        logger.info(f"PDF 加载完成: {file_path}, 共 {len(documents)} 页")
    except Exception as e:
        logger.error(f"PDF 加载失败: {file_path}, 错误: {e}")
        raise
    return documents


def load_markdown(file_path: str) -> list[Document]:
    """加载 Markdown 文件"""
    documents: list[Document] = []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read().strip()
        if text:
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": os.path.basename(file_path),
                        "type": "markdown",
                    },
                )
            )
        logger.info(f"Markdown 加载完成: {file_path}")
    except Exception as e:
        logger.error(f"Markdown 加载失败: {file_path}, 错误: {e}")
        raise
    return documents


def load_txt(file_path: str) -> list[Document]:
    """加载纯文本文件"""
    documents: list[Document] = []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read().strip()
        if text:
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": os.path.basename(file_path),
                        "type": "txt",
                    },
                )
            )
        logger.info(f"TXT 加载完成: {file_path}")
    except Exception as e:
        logger.error(f"TXT 加载失败: {file_path}, 错误: {e}")
        raise
    return documents


def load_docx(file_path: str) -> list[Document]:
    """加载 Word (.docx) 文件"""
    from docx import Document as DocxDocument

    documents: list[Document] = []
    try:
        doc = DocxDocument(file_path)
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        if paragraphs:
            full_text = "\n\n".join(paragraphs)
            documents.append(
                Document(
                    page_content=full_text,
                    metadata={
                        "source": os.path.basename(file_path),
                        "type": "docx",
                    },
                )
            )
        logger.info(f"DOCX 加载完成: {file_path}, 段落数: {len(paragraphs)}")
    except Exception as e:
        logger.error(f"DOCX 加载失败: {file_path}, 错误: {e}")
        raise
    return documents


def load_document(file_path: str) -> list[Document]:
    """统一文档加载接口，根据文件扩展名自动选择加载器"""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的文件格式: {ext}，支持: {SUPPORTED_EXTENSIONS}")

    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext in (".md", ".markdown"):
        return load_markdown(file_path)
    elif ext == ".txt":
        return load_txt(file_path)
    elif ext == ".docx":
        return load_docx(file_path)

    return []
